"""Generate presentation_for_GK.docx from markdown content."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# ── Title ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ДОКЛАД')
r.bold = True
r.font.size = Pt(18)
r.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Канбан-доска: управление задачами в «Союз-PLM»')
r.bold = True
r.font.size = Pt(15)
r.font.name = 'Times New Roman'

doc.add_paragraph()

# ── Helper functions ──
def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14) if level == 1 else Pt(13)
    r.font.name = 'Times New Roman'
    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return p

def add_para(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    return p

def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(text)
    r.italic = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
    # Rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            r = cell.paragraphs[0].add_run(val)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
    doc.add_paragraph()  # spacing after table

# ═══════════════════════════════════════════════════
# CONTENT
# ═══════════════════════════════════════════════════

add_heading('1. ИДЕЯ — ОТКУДА КАНБАН')

add_para('В 1950-х инженер Toyota Тайити Оно решил простую задачу: как сделать так, чтобы каждый участок производства точно знал, что и когда делать — без совещаний, без обходов цехов, без потерянных поручений.')

add_para('Решение оказалось элементарным: карточки на доске. Каждая задача — карточка. Колонки — этапы работы. Взял задачу — перенёс карточку. Сделал — передвинул дальше. Руководитель смотрит на доску и мгновенно видит всю картину: что стоит, что горит, кто перегружен.')

add_para('Этот метод — канбан — сделал Toyota самым эффективным автопроизводителем в мире. Сегодня его используют все: от авиастроения до NASA.')

add_para('Аналогия с нашим КБ прямая: ', bold_prefix='')
p = doc.paragraphs[-1]
p.clear()
r = p.add_run('Аналогия с нашим КБ прямая: ')
r.bold = True
r.font.name = 'Times New Roman'
r.font.size = Pt(13)
r = p.add_run('у Toyota — детали движутся по участкам, у нас — задачи движутся по этапам. У Toyota — карточка сигналит «нужна деталь», у нас — карточка сигналит «нужно сделать». Принцип один: визуальное управление потоком работ.')
r.font.name = 'Times New Roman'
r.font.size = Pt(13)

# ───────────────────────────────────────────────────
add_heading('2. ПРОБЛЕМА')

add_para('Сегодня задачи в подразделениях живут в головах, в почте, в устных поручениях.')

add_bullet('руководитель не видит загрузку сотрудников, пока не обойдёт кабинеты', bold_prefix='Непрозрачно: ')
add_bullet('устное поручение забылось — узнали, когда сорвали срок', bold_prefix='Теряется: ')
add_bullet('всё «срочно», непонятно, что делать первым', bold_prefix='Нет приоритетов: ')
add_bullet('отпуск, больничный — контекст теряется, работа встаёт', bold_prefix='Передача дел: ')

add_quote('Мы тратим время не на работу, а на выяснение — кто что делает и на каком этапе.')

# ───────────────────────────────────────────────────
add_heading('3. РЕШЕНИЕ — КАНБАН-ДОСКА В «СОЮЗ»')

add_para('Электронная канбан-доска прямо внутри «Союз-PLM». Не нужно осваивать новую программу — открывается как ещё один экран в привычной среде.')

add_para('Четыре колонки — четыре состояния работы:')

add_table(
    ['Надо сделать', 'В работе', 'Ожидание', 'Готово'],
    [['Новые задачи', 'Выполняются сейчас', 'Ждём ответ / согласование', 'Завершены']]
)

add_para('Каждая задача — карточка с названием, описанием, приоритетом, сроком и исполнителем. Перетащил в другую колонку — статус обновился.')

# ───────────────────────────────────────────────────
add_heading('4. ВОЗМОЖНОСТИ')

add_para('Управление задачами:', bold_prefix='')
doc.paragraphs[-1].runs[0].bold = True
add_bullet('Создание задачи себе, конкретному сотруднику или сразу группе')
add_bullet('Четыре уровня приоритета — сразу видно, что горит')
add_bullet('Сроки — просроченные задачи подсвечиваются красным')

add_para('Прозрачность для руководства:', bold_prefix='')
doc.paragraphs[-1].runs[0].bold = True
add_bullet('Начальник сектора видит все задачи своих сотрудников на одной доске')
add_bullet('Начальник отделения видит картину по всем секторам')
add_bullet('Формирование отчётов за любой период')

add_para('Работа с документами:', bold_prefix='')
doc.paragraphs[-1].runs[0].bold = True
add_bullet('К задаче прикрепляются документы прямо из «Союза» — чертежи, модели, расчёты')
add_bullet('Весь контекст привязан к задаче, а не разбросан по переписке')

add_para('Обсуждение и контроль:', bold_prefix='')
doc.paragraphs[-1].runs[0].bold = True
add_bullet('Комментарии внутри карточки — обсуждение в одном месте')
add_bullet('История изменений: кто и когда менял статус, переназначал')
add_bullet('Уведомления при назначении новой задачи')

# ───────────────────────────────────────────────────
add_heading('5. ЧТО ЭТО ДАЁТ')

add_table(
    ['Было', 'Стало'],
    [
        ['«Ты сделал?» — «Какое именно?»', 'Открыл доску — всё видно'],
        ['Задачи в почте, блокноте, устно', 'Единое место для всех задач'],
        ['Руководитель обходит всех лично', 'Смотрит доску — видит загрузку и статусы'],
        ['Совещание: «расскажите кто чем занят»', 'Доска уже показывает — совещание короче'],
        ['Поручение по телефону — забылось', 'Задача на доске — не потеряется'],
    ]
)

add_quote('Обновить статус карточки — 5 секунд. Найти потерянную задачу без доски — полдня.')

# ───────────────────────────────────────────────────
add_heading('6. ПЕРСПЕКТИВЫ')

add_bullet('задачи из годового/квартального плана попадают на доску', bold_prefix='Связь с планированием — ')
add_bullet('согласование, выпуск КД → задача на доске', bold_prefix='Связь с рабочими процессами — ')
add_bullet('подключение любого подразделения', bold_prefix='Масштабирование — ')

# ───────────────────────────────────────────────────
add_heading('7. ПРЕДЛОЖЕНИЕ')

add_para('Предлагаю утвердить запуск опытной эксплуатации канбан-доски в рамках подразделения. По результатам — доклад с обратной связью.')

add_quote('Тот же принцип, что вывел Toyota в лидеры мирового производства — только адаптированный под наше КБ и встроенный в систему, которой мы уже пользуемся.')

# ═══════════════════════════════════════════════════
out = r'c:\MYPROJECTS\kanbanSoyz_Actual\presentation_for_GK.docx'
doc.save(out)
print('OK:', out)
